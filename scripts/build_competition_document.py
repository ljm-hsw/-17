#!/usr/bin/env python3
"""Build the competition DOCX from the retained official template.

The reference template is never modified. All generated graphics and the final
DOCX are deterministic outputs under docs/competition/.
"""

from __future__ import annotations

import hashlib
import math
import re
import shutil
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "docs" / "competition"
ASSET_DIR = DOC_DIR / "assets"
REFERENCE = DOC_DIR / "reference" / "2026年全国大学生物联网设计竞赛设计文档模板.docx"
MANUSCRIPT = DOC_DIR / "游迹织梦-作品设计文档-正文.md"
FINAL = DOC_DIR / "游迹织梦-作品设计文档.docx"
REFERENCE_SHA256 = "b639443ed7bfd8e3875e5dee962c74ecedc600e550125bc8c5b4b725fe046047"

SONG_FONT = "/System/Library/Fonts/Supplemental/Songti.ttc"
HEI_FONT = "/System/Library/Fonts/STHeiti Medium.ttc"

INK = "263238"
GREEN = "2F8F72"
LIGHT_GREEN = "E8F3EE"
ORANGE = "E9852C"
LIGHT_ORANGE = "FFF1E4"
BLUE = "3978A8"
LIGHT_BLUE = "E8F1F8"
GRAY = "66737A"
LIGHT_GRAY = "F4F6F5"
WHITE = "FFFFFF"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def rgb(value: str) -> tuple[int, int, int]:
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(HEI_FONT if bold else SONG_FONT, size=size, index=0)


def rounded_box(draw, xy, fill, outline, radius=24, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, xy, text, fnt, fill=rgb(INK), spacing=8):
    draw.multiline_text(xy, text, font=fnt, fill=fill, anchor="mm", align="center", spacing=spacing)


def arrow(draw, start, end, color=rgb(GREEN), width=8):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 22
    left = (end[0] - size * math.cos(angle - 0.55), end[1] - size * math.sin(angle - 0.55))
    right = (end[0] - size * math.cos(angle + 0.55), end[1] - size * math.sin(angle + 0.55))
    draw.polygon([end, left, right], fill=color)


def canvas(title: str, subtitle: str = ""):
    image = Image.new("RGB", (1800, 1000), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1800, 110), fill=rgb(INK))
    draw.text((70, 48), title, font=font(42, True), fill=rgb(WHITE), anchor="lm")
    if subtitle:
        draw.text((1730, 54), subtitle, font=font(24), fill=(212, 224, 221), anchor="rm")
    return image, draw


def save_image(image: Image.Image, name: str):
    path = ASSET_DIR / name
    image.save(path, "PNG", optimize=True)
    return path


def build_journey():
    image, draw = canvas("游客完整旅程", "从进入小程序到形成文化记忆")
    labels = [
        ("01", "微信登录\n知情告知", GREEN),
        ("02", "绑定手环\nNFC / 激活码", BLUE),
        ("03", "选择路线\n开始游览", ORANGE),
        ("04", "实体刷卡\n可信事件", GREEN),
        ("05", "地图进度\n知识与推荐", BLUE),
        ("06", "确认拍照\n事件关联", ORANGE),
        ("07", "AI 游记\n保存与分享", GREEN),
    ]
    xs = [150, 400, 650, 900, 1150, 1400, 1650]
    for idx, ((num, label, color), x) in enumerate(zip(labels, xs)):
        if idx < len(xs) - 1:
            arrow(draw, (x + 78, 510), (xs[idx + 1] - 78, 510), rgb(GRAY), 5)
        draw.ellipse((x - 74, 300, x + 74, 448), fill=rgb(color), outline=rgb(INK), width=3)
        centered(draw, (x, 374), num, font(38, True), rgb(WHITE))
        rounded_box(draw, (x - 110, 500, x + 110, 690), rgb("F9FAF9"), rgb(color), 18, 4)
        centered(draw, (x, 595), label, font(27, True), rgb(INK), 12)
    centered(draw, (900, 830), "主动交互 · 清晰反馈 · 全程可回顾 · 权限可撤回", font(34, True), rgb(GREEN))
    save_image(image, "fig01_journey.png")


def build_value_loop():
    image, draw = canvas("“感知—可信—智能—记忆”价值闭环", "同一事件沿数据链持续增值")
    nodes = [
        ((330, 290), "真实感知", "ESP32-S3\nRFID/NFC 刷卡", GREEN),
        ((1270, 290), "可信数据", "认证 · 幂等\n绑定 · 会话", BLUE),
        ((1270, 720), "智能服务", "地图 · 推荐\n问答 · 管理", ORANGE),
        ((330, 720), "文化记忆", "路线回顾\n照片 · 游记", GREEN),
    ]
    for (x, y), title, detail, color in nodes:
        rounded_box(draw, (x - 230, y - 120, x + 230, y + 120), rgb(WHITE), rgb(color), 32, 7)
        centered(draw, (x, y - 35), title, font(38, True), rgb(color))
        centered(draw, (x, y + 48), detail, font(26), rgb(INK))
    arrow(draw, (565, 290), (1035, 290), rgb(BLUE), 9)
    arrow(draw, (1270, 415), (1270, 590), rgb(ORANGE), 9)
    arrow(draw, (1035, 720), (565, 720), rgb(GREEN), 9)
    arrow(draw, (330, 590), (330, 415), rgb(GREEN), 9)
    rounded_box(draw, (700, 400, 1100, 610), rgb(LIGHT_GRAY), rgb(INK), 90, 4)
    centered(draw, (900, 475), "统一事件事实", font(36, True), rgb(INK))
    centered(draw, (900, 545), "CheckinEvent", font(28), rgb(GRAY))
    save_image(image, "fig02_value_loop.png")


def build_agent():
    image, draw = canvas("AI Agent 受控工具与事实校验工作流", "模型不直接访问数据库")
    items = [
        ("用户请求", 120, GREEN),
        ("身份与权限", 360, BLUE),
        ("个人事实工具", 600, GREEN),
        ("审核知识检索", 840, BLUE),
        ("Dify 生成", 1080, ORANGE),
        ("事实校验", 1320, GREEN),
        ("结果与版本", 1560, BLUE),
    ]
    for idx, (label, x, color) in enumerate(items):
        rounded_box(draw, (x - 95, 330, x + 95, 650), rgb(WHITE), rgb(color), 28, 5)
        centered(draw, (x, 490), label, font(25, True), rgb(INK), 6)
        draw.ellipse((x - 34, 235, x + 34, 303), fill=rgb(color))
        centered(draw, (x, 269), str(idx + 1), font(25, True), rgb(WHITE))
        if idx < len(items) - 1:
            arrow(draw, (x + 96, 490), (items[idx + 1][1] - 98, 490), rgb(GRAY), 5)
    rounded_box(draw, (530, 760, 1270, 900), rgb(LIGHT_ORANGE), rgb(ORANGE), 22, 3)
    centered(draw, (900, 830), "发送给模型：点位名称、顺序、时间、公开知识与授权状态\n不发送：原始 UID、微信标识、设备密钥、未授权照片", font(27, True), rgb(INK), 10)
    save_image(image, "fig03_agent.png")


def build_sequence():
    image, draw = canvas("RFID/NFC 打卡与拍照关联时序", "拍照点：授权 + 有效打卡 + 即时提示 + 用户确认")
    actors = [(180, "游客"), (500, "ESP32-S3\n点位终端"), (850, "Django\n后端"), (1180, "微信小程序"), (1530, "拍照终端")]
    for x, name in actors:
        rounded_box(draw, (x - 105, 150, x + 105, 235), rgb(LIGHT_GREEN), rgb(GREEN), 18, 3)
        centered(draw, (x, 193), name, font(25, True), rgb(INK))
        draw.line((x, 235, x, 900), fill=rgb("B5C0BC"), width=3)
    events = [
        (300, 180, 500, "刷卡"),
        (390, 500, 850, "签名上传候选事件"),
        (480, 850, 500, "accepted / event_id"),
        (570, 850, 1180, "推送打卡结果与拍照提示"),
        (660, 1180, 180, "展示说明，用户确认"),
        (750, 1180, 850, "申请短时拍照令牌"),
        (840, 850, 1530, "令牌绑定 event_id，单次触发"),
    ]
    for y, x1, x2, label in events:
        color = rgb(ORANGE) if "拍照" in label or "确认" in label or "令牌" in label else rgb(BLUE)
        arrow(draw, (x1, y), (x2, y), color, 5)
        draw.text(((x1 + x2) / 2, y - 17), label, font=font(22, True), fill=rgb(INK), anchor="ms")
    save_image(image, "fig04_sequence.png")


def build_miniprogram():
    source = Image.open(ASSET_DIR / "miniprogram-reference-1.jpg").convert("RGB")
    source.thumbnail((1680, 790), Image.Resampling.LANCZOS)
    image, draw = canvas("微信小程序关键页面设计", "高保真界面设计，不作为运行截图")
    x = (1800 - source.width) // 2
    image.paste(source, (x, 140))
    draw.rectangle((0, 930, 1800, 1000), fill=rgb(LIGHT_ORANGE))
    centered(draw, (900, 965), "登录绑定 · 首页进度 · 校园地图 · 点位详情", font(28, True), rgb(INK))
    save_image(image, "fig05_miniprogram.png")


def build_admin():
    image, draw = canvas("Web 运营管理系统界面", "比赛展示与完整运营管理共用")
    rounded_box(draw, (70, 150, 1730, 920), rgb("F7F9F8"), rgb("CDD5D1"), 24, 3)
    draw.rectangle((70, 150, 340, 920), fill=rgb(INK))
    draw.text((110, 205), "游迹织梦", font=font(34, True), fill=rgb(WHITE))
    menus = ["运营看板", "用户管理", "场景与点位", "卡片与绑定", "设备状态", "打卡追溯", "审计日志"]
    for i, menu in enumerate(menus):
        y = 300 + i * 76
        if i == 0:
            rounded_box(draw, (95, y - 25, 315, y + 35), rgb(GREEN), rgb(GREEN), 12, 2)
        draw.text((125, y), menu, font=font(24, i == 0), fill=rgb(WHITE), anchor="lm")
    draw.text((390, 210), "江安校区运营总览", font=font(38, True), fill=rgb(INK))
    cards = [("今日有效打卡", "128", GREEN), ("活跃游客", "46", BLUE), ("已绑定卡片", "83", ORANGE), ("在线设备", "7 / 8", GREEN)]
    for i, (title, value, color) in enumerate(cards):
        x1 = 390 + i * 315
        rounded_box(draw, (x1, 270, x1 + 280, 430), rgb(WHITE), rgb(color), 18, 3)
        draw.text((x1 + 25, 310), title, font=font(22), fill=rgb(GRAY))
        draw.text((x1 + 25, 378), value, font=font(42, True), fill=rgb(color), anchor="lm")
    rounded_box(draw, (390, 480, 1060, 850), rgb(WHITE), rgb("D4DAD7"), 18, 2)
    draw.text((425, 525), "近 7 日打卡趋势", font=font(26, True), fill=rgb(INK))
    points = [(445, 780), (540, 700), (635, 735), (730, 620), (825, 650), (920, 540), (1015, 585)]
    draw.line(points, fill=rgb(GREEN), width=8, joint="curve")
    for x, y in points:
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=rgb(GREEN))
    rounded_box(draw, (1090, 480, 1685, 850), rgb(WHITE), rgb("D4DAD7"), 18, 2)
    draw.text((1125, 525), "热门点位", font=font(26, True), fill=rgb(INK))
    spots = [("江安图书馆", 92), ("明远湖", 80), ("长桥", 64), ("知识广场", 48)]
    for i, (name, value) in enumerate(spots):
        y = 600 + i * 62
        draw.text((1125, y), name, font=font(21), fill=rgb(INK), anchor="lm")
        draw.rounded_rectangle((1300, y - 13, 1610, y + 13), radius=13, fill=rgb("E6ECE9"))
        draw.rounded_rectangle((1300, y - 13, 1300 + int(310 * value / 100), y + 13), radius=13, fill=rgb(BLUE))
    save_image(image, "fig06_admin.png")


def build_architecture():
    image, draw = canvas("四端协同总体架构与信任边界", "后端是唯一业务事实源")
    columns = [
        (100, 360, "设备区", GREEN, ["ESP32-S3", "RFID/NFC", "LED / 蜂鸣器", "可选拍照终端"]),
        (500, 760, "用户区", BLUE, ["微信小程序", "登录与绑卡", "地图与记录", "AI 导游"]),
        (1040, 1300, "运营区", ORANGE, ["Vue 3 管理端", "点位与设备", "卡片与用户", "审计追溯"]),
        (1440, 1700, "AI 区", GREEN, ["Dify 工作流", "云端大模型", "知识检索", "事实校验"]),
    ]
    for x1, x2, title, color, items in columns:
        rounded_box(draw, (x1, 185, x2, 830), rgb(WHITE), rgb(color), 28, 5)
        draw.rectangle((x1, 185, x2, 280), fill=rgb(color))
        centered(draw, ((x1 + x2) / 2, 232), title, font(31, True), rgb(WHITE))
        for i, item in enumerate(items):
            rounded_box(draw, (x1 + 35, 330 + i * 105, x2 - 35, 400 + i * 105), rgb(LIGHT_GRAY), rgb("CBD3CF"), 14, 2)
            centered(draw, ((x1 + x2) / 2, 365 + i * 105), item, font(23, True), rgb(INK))
    rounded_box(draw, (770, 180, 1030, 835), rgb(INK), rgb(INK), 28, 3)
    centered(draw, (900, 265), "Django API", font(34, True), rgb(WHITE))
    for i, item in enumerate(["accounts", "scenes", "iot", "visits", "media", "ai"]):
        centered(draw, (900, 370 + i * 70), item, font(24, True), (221, 234, 229))
    for x in [360, 760, 1040, 1440]:
        if x < 900:
            arrow(draw, (x, 880), (760, 880), rgb(GRAY), 4)
        else:
            arrow(draw, (1040, 880), (x, 880), rgb(GRAY), 4)
    centered(draw, (900, 920), "PostgreSQL · 对象存储 · 审计日志", font(28, True), rgb(GRAY))
    save_image(image, "fig07_architecture.png")


def build_erd():
    image, draw = canvas("核心数据实体关系", "事件连接用户、实体媒介、空间与设备")
    nodes = {
        "User": (180, 250, GREEN), "Card": (180, 520, BLUE), "CardBinding": (520, 385, ORANGE),
        "VisitSession": (880, 250, GREEN), "CheckinEvent": (880, 520, INK), "Scene": (1250, 210, BLUE),
        "Spot": (1250, 480, GREEN), "Device": (1580, 480, ORANGE), "PhotoAsset": (1250, 750, BLUE),
    }
    for name, (x, y, color) in nodes.items():
        rounded_box(draw, (x - 130, y - 65, x + 130, y + 65), rgb(WHITE), rgb(color), 18, 5)
        centered(draw, (x, y), name, font(26, True), rgb(color if color != INK else WHITE) if False else rgb(INK))
    links = [
        ("User", "CardBinding", "1:N"), ("Card", "CardBinding", "1:N 历史"),
        ("User", "VisitSession", "1:N"), ("CardBinding", "CheckinEvent", "发生时绑定"),
        ("VisitSession", "CheckinEvent", "1:N"), ("Scene", "Spot", "1:N"),
        ("Spot", "CheckinEvent", "1:N"), ("Device", "CheckinEvent", "1:N"),
        ("CheckinEvent", "PhotoAsset", "0:1"),
    ]
    for a, b, label in links:
        x1, y1, _ = nodes[a]; x2, y2, _ = nodes[b]
        arrow(draw, (x1, y1), (x2, y2), rgb("9AA6A1"), 4)
        draw.text(((x1 + x2) / 2, (y1 + y2) / 2 - 12), label, font=font(18, True), fill=rgb(GRAY), anchor="ms")
    rounded_box(draw, (420, 790, 850, 900), rgb(LIGHT_ORANGE), rgb(ORANGE), 18, 3)
    centered(draw, (635, 845), "唯一约束：卡片当前归属\n用户主卡 · 自然日会话 · event_id", font(23, True), rgb(INK))
    save_image(image, "fig08_erd.png")


def build_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_journey()
    build_value_loop()
    build_agent()
    build_sequence()
    build_miniprogram()
    build_admin()
    build_architecture()
    build_erd()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, east_asia="Arial Unicode MS", latin="Arial Unicode MS", size=None, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, east_asia, latin, size, bold=False, color=INK):
    style.font.name = latin
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def apply_template_derived_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Arial Unicode MS", "Arial Unicode MS", 10.5, False, INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.25
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(21)

    title = doc.styles["Title"]
    set_style_font(title, "Arial Unicode MS", "Arial Unicode MS", 26, True, INK)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, before, after in (("Heading 1", 18, 18, 18), ("Heading 2", 15, 14, 8), ("Heading 3", 12, 10, 4)):
        style = doc.styles[name]
        set_style_font(style, "Arial Unicode MS", "Arial Unicode MS", size, True, INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Pt(0)
        if name == "Heading 1":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def clear_document_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_page_geometry(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.98)
    section.top_margin = Inches(1.05)
    section.bottom_margin = Inches(0.95)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def set_page_numbering(section, fmt=None, start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    elif qn("w:start") in pg_num.attrib:
        del pg_num.attrib[qn("w:start")]


def add_field(paragraph, instruction, cached=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    if cached:
        cache_run = paragraph.add_run(cached)
        set_run_font(cache_run, size=10.5)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_page_footer(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, " PAGE ", "1")


def add_header(section):
    section.header.is_linked_to_previous = False
    header = section.header
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("全国大学生物联网设计竞赛设计文档  [游迹织梦]")
    set_run_font(r, size=9, color=GRAY)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), INK)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_new_section(doc: Document, *, fmt=None, start=None, header=True):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(section)
    if header:
        add_header(section)
        add_page_footer(section)
    set_page_numbering(section, fmt, start)
    return section


def add_cover(doc: Document):
    section = doc.sections[0]
    set_page_geometry(section)
    section.top_margin = Inches(1.05)
    section.bottom_margin = Inches(0.75)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("全国大学生物联网设计竞赛")
    set_run_font(r, size=26, bold=True)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run("游迹织梦")
    set_run_font(r, size=36, bold=True, color=GREEN)

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["学校名称：", "团队名称：", "", "队长：", "队员1：", "队员2：", "队员3："]
    values = ["", "成都潮人", "", "", "", "", ""]
    for i, row in enumerate(table.rows):
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(9.2)
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 130, 160, 130, 160)
            cell.text = labels[i] if j == 0 else values[i]
            for p2 in cell.paragraphs:
                p2.paragraph_format.first_line_indent = Pt(0)
                p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p2.runs:
                    set_run_font(run, size=13, bold=(j == 0))
        if i == 2:
            row.height = Cm(0.35)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("全国大学生物联网设计竞赛组委会")
    set_run_font(r, size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年7月")
    set_run_font(r, size=14)


def add_abstract(doc: Document, text: str, keywords: str):
    add_new_section(doc, fmt="upperRoman", start=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("游迹织梦")
    set_run_font(r, size=20, bold=True)
    doc.add_paragraph("摘  要", style="Heading 1")
    p = doc.add_paragraph(text.strip())
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Pt(21)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        set_run_font(r, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.first_line_indent = Pt(0)
    label = p.add_run("关键词：")
    set_run_font(label, size=10.5, bold=True)
    value = p.add_run(keywords)
    set_run_font(value, size=10.5)


def add_toc(doc: Document, entries):
    add_new_section(doc, fmt="upperRoman")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("目  录")
    set_run_font(r, size=18, bold=True)
    for level, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm((level - 1) * 0.75)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        r = p.add_run(title)
        set_run_font(r, size=10 if level > 1 else 10.5, bold=(level == 1))
        p.add_run("\t")
        page = p.add_run(" ")
        set_run_font(page, size=10)


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("**", "")


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    # Long API paths and bibliography entries contain unbreakable Latin tokens.
    # Left alignment avoids Word stretching a short Chinese prefix across a line.
    if text.startswith("主要用户接口包括") or re.match(r"^\[[0-9]+\]", text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    text = clean_inline(text)
    for part in re.split(r"(\[[0-9]+\])", text):
        run = p.add_run(part)
        set_run_font(run, size=10.5)
        if re.fullmatch(r"\[[0-9]+\]", part):
            run.font.superscript = True
    return p


def add_picture(doc: Document, alt: str, rel_path: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    shape = run.add_picture(str(DOC_DIR / rel_path), width=Inches(6.25))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run(alt)
    set_run_font(r, size=9.5, bold=False)


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for j, value in enumerate(rows[0]):
        table.rows[0].cells[j].text = clean_inline(value)
    for source_row in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(source_row):
            cells[j].text = clean_inline(value)
    total_width = Cm(15.6)
    weights = []
    for col in range(len(rows[0])):
        max_len = max(len(row[col]) for row in rows)
        weights.append(max(8, min(max_len, 34)))
    weight_sum = sum(weights)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = total_width * weights[j] / weight_sum
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 90, 110, 90, 110)
            if i == 0:
                set_cell_shading(cell, INK)
            elif i % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=8.7, bold=(i == 0), color=(WHITE if i == 0 else INK))
    set_repeat_table_header(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_manuscript(text: str):
    lines = text.splitlines()
    abstract = text.split("## 摘要", 1)[1].split("**关键词：**", 1)[0].strip()
    keywords = text.split("**关键词：**", 1)[1].split("# 第一章", 1)[0].strip()
    entries = []
    for line in lines:
        if line.startswith("# ") and line != "# 游迹织梦":
            entries.append((1, line[2:].strip()))
        elif line.startswith("## ") and line != "## 摘要":
            entries.append((2, line[3:].strip()))
        elif line.startswith("### "):
            entries.append((3, line[4:].strip()))
    return lines, abstract, keywords, entries


def build_docx():
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("Official reference template hash mismatch")
    shutil.copy2(REFERENCE, FINAL)
    doc = Document(FINAL)
    clear_document_body(doc)
    apply_template_derived_styles(doc)
    add_cover(doc)

    text = MANUSCRIPT.read_text(encoding="utf-8")
    lines, abstract, keywords, entries = parse_manuscript(text)
    add_abstract(doc, abstract, keywords)
    add_toc(doc, [(1, "摘  要")] + [entry for entry in entries if entry[0] <= 2])

    in_body = False
    pending_table = []
    current_chapter = 0
    for raw in lines:
        line = raw.strip()
        if line == "# 第一章 设计需求分析":
            in_body = True
        if not in_body or not line:
            continue
        if line.startswith("|"):
            if re.match(r"^\|[-: |]+\|$", line):
                continue
            pending_table.append([x.strip() for x in line.strip("|").split("|")])
            continue
        if pending_table:
            add_table(doc, pending_table)
            pending_table = []
        if line.startswith("# "):
            current_chapter += 1
            add_new_section(doc, fmt="decimal", start=1 if current_chapter == 1 else None)
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("!["):
            match = re.match(r"!\[([^]]+)\]\(([^)]+)\)", line)
            if match:
                add_picture(doc, match.group(1), match.group(2))
        elif line.startswith("[") and current_chapter == 6:
            p = add_body_paragraph(doc, line)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.left_indent = Pt(18)
        else:
            add_body_paragraph(doc, line)
    if pending_table:
        add_table(doc, pending_table)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    core = doc.core_properties
    core.title = "游迹织梦"
    core.subject = "2026年全国大学生物联网设计竞赛乐鑫科技赛道作品设计文档"
    core.author = "成都潮人"
    core.keywords = "智慧文旅, ESP32-S3, RFID/NFC, 微信小程序, 生成式人工智能"
    core.comments = "基于官方模板制作"

    # Mark every first row as a semantic/repeating header. This includes any
    # table inherited from the official template package as well as tables
    # created from the manuscript.
    for table in doc.tables:
        if table.rows:
            set_repeat_table_header(table.rows[0])

    doc.save(FINAL)
    if sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("Reference template changed during build")


def normalized(value: str) -> str:
    return re.sub(r"\s+", "", value).replace("—", "-").replace("–", "-")


def patch_toc_from_pdf_text(pdf_text: Path):
    """Fill visible TOC page numbers from a rendered PDF text extraction."""
    text = MANUSCRIPT.read_text(encoding="utf-8")
    _, _, _, entries = parse_manuscript(text)
    entries = [entry for entry in entries if entry[0] <= 2]
    pages = pdf_text.read_text(encoding="utf-8", errors="replace").split("\f")
    page_norm = [normalized(page) for page in pages]
    chapter_one = normalized("第一章 设计需求分析")
    body_matches = [i for i, page in enumerate(page_norm) if chapter_one in page]
    try:
        body_start = body_matches[-1]
    except IndexError as exc:
        raise RuntimeError("Cannot locate chapter one in rendered PDF text") from exc

    mapping = {"摘  要": "I"}
    for _, title in entries:
        target = normalized(title)
        found = None
        for page_index in range(body_start, len(page_norm)):
            if target in page_norm[page_index]:
                found = page_index
                break
        if found is None:
            raise RuntimeError(f"Cannot locate TOC heading in rendered PDF: {title}")
        mapping[title] = str(found - body_start + 1)

    doc = Document(FINAL)
    in_toc = False
    patched = 0
    for paragraph in doc.paragraphs:
        text_value = paragraph.text.strip()
        if text_value == "目  录":
            in_toc = True
            continue
        if in_toc and paragraph.style and paragraph.style.name == "Heading 1":
            break
        if not in_toc or "\t" not in paragraph.text:
            continue
        title = paragraph.text.split("\t", 1)[0].strip()
        if title not in mapping:
            continue
        paragraph.runs[-1].text = mapping[title]
        set_run_font(paragraph.runs[-1], size=10)
        patched += 1
    if patched != len(mapping):
        raise RuntimeError(f"TOC patch incomplete: {patched}/{len(mapping)}")
    doc.save(FINAL)
    print(f"toc_patched={patched} body_start_physical_page={body_start + 1}")


def main():
    if not REFERENCE.exists() or not MANUSCRIPT.exists():
        raise SystemExit("Missing reference template or manuscript")
    if len(sys.argv) == 3 and sys.argv[1] == "--patch-toc":
        patch_toc_from_pdf_text(Path(sys.argv[2]))
        return
    build_assets()
    build_docx()
    print(FINAL)
    print(f"size={FINAL.stat().st_size} sha256={sha256(FINAL)}")


if __name__ == "__main__":
    main()
